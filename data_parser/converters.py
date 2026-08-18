import asyncio
import logging
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import aiofiles

logger = logging.getLogger(__name__)


class FileConverter:
    @staticmethod
    def _extract_pdf_text(pdf_path: str) -> str:
        reader = PdfReader(pdf_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)

    @staticmethod
    def _extract_docx_text(docx_path: str) -> str:
        doc = Document(docx_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    @classmethod
    async def pdf_to_txt_async(cls, pdf_path: str) -> str:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, cls._extract_pdf_text, pdf_path)

    @classmethod
    async def docx_to_txt_async(cls, docx_path: str) -> str:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, cls._extract_docx_text, docx_path)

    @classmethod
    async def convert_file_async(cls, input_path: str, output_dir: str) -> Path:
        input_path = Path(input_path)
        ext = input_path.suffix.lower()

        if ext == ".pdf":
            text = await cls.pdf_to_txt_async(str(input_path))
        elif ext == ".docx":
            text = await cls.docx_to_txt_async(str(input_path))
        else:
            raise ValueError(f"Неподдерживаемый формат: {ext}")

        out_path = Path(output_dir) / (input_path.stem + ".txt")
        async with aiofiles.open(out_path, "w", encoding="utf-8") as f:
            await f.write(text)

        return out_path
